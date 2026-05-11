#!/usr/bin/env python3
"""
Word 文档格式化工具 — 公文格式规范版（v4 递归区间层级识别）
========================================================
根据党政机关公文格式规范，自动格式化 Word 文档。

格式规范：
  • 标题：     方正小标宋_GBK 二号（22pt），居中
  • 一级标题： 黑体 三号（16pt）
  • 二级标题： 楷体_GB2312 三号（16pt）
  • 三级标题： 仿宋_GB2312 三号（16pt）
  • 四级标题： 仿宋_GB2312 三号（16pt）
• 正文：     仿宋_GB2312 三号（16pt），首行缩进2字符，行距固定28磅

v4 改进：
  - 递归区间划分：先全局找 h1，再在每个 h1 之间独立找 h2，递归类推
  - 无编号标题 = 当前递归层的格式
  - 每个同级标题之间的区间独立遍历
  - 自动编号 → 纯文本（兼容 Word 自动编号列表）
  - 按规范重写编号：一、→（一）→ 1. →（1）
  - 移除正文空行
  - 表格统一样式：单实线边框，表头灰色

使用方法：
    format_word.exe <输入文件.docx> [输出文件.docx]

    若不指定输出文件，默认在输入文件同目录生成 "格式化_<原文件名>.docx"。
"""

import re
import os
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════
# 配置区 — 字体文件
# ═══════════════════════════════════════════════════════════════════
FONT_DIR = os.path.dirname(os.path.abspath(__file__))

FONT_FILES = {
    '方正小标宋_GBK': '方正小标宋_GBK.TTF',
    '仿宋_GB2312':    '仿宋_GB2312.TTF',
    '黑体':           '黑体_GB18030.TTF',
    '楷体_GB2312':    '楷体_GB2312.TTF',
}

FONT_NAMES = {
    'title': '方正小标宋_GBK',
    'h1':    '黑体',
    'h2':    '楷体_GB2312',
    'h3':    '仿宋_GB2312',
    'h4':    '仿宋_GB2312',
    'body':  '仿宋_GB2312',
}

# ═══════════════════════════════════════════════════════════════════
# 配置区 — 字号与格式参数
# ═══════════════════════════════════════════════════════════════════

FONT_SIZES = {
    'title': Pt(22),
    'h1':    Pt(16),
    'h2':    Pt(16),
    'h3':    Pt(16),
    'h4':    Pt(16),
    'body':  Pt(16),
}

FONT_BOLD = {
    'title': False,
    'h1':    False,
    'h2':    False,
    'h3':    False,
    'h4':    False,
    'body':  False,
}

ALIGNMENT = {
    'title': WD_ALIGN_PARAGRAPH.CENTER,
    'h1':    WD_ALIGN_PARAGRAPH.LEFT,
    'h2':    WD_ALIGN_PARAGRAPH.LEFT,
    'h3':    WD_ALIGN_PARAGRAPH.LEFT,
    'h4':    WD_ALIGN_PARAGRAPH.LEFT,
    'body':  WD_ALIGN_PARAGRAPH.JUSTIFY,
}

LINE_SPACING_FIXED = Pt(28)
FIRST_LINE_INDENT = Pt(32)
MAX_TITLE_LENGTH = 60

LEVEL_NAMES = ['h1', 'h2', 'h3', 'h4']


# ═══════════════════════════════════════════════════════════════════
# 编号风格分类
# ═══════════════════════════════════════════════════════════════════

def classify_numbering(text):
    """
    提取段落文本的编号风格。

    返回 (style_key, prefix_text)：
      style_key  — 用于聚类同层标题的标识
      prefix_text — 编号前缀原文
    无编号时返回 (None, None)。
    """
    # 多级编号 1.1, 2.3.1（需在 arabic_dot 之前检查）
    m = re.match(r'^(\d+\.\d+)', text)
    if m:
        return ('arabic_multi', m.group(1))

    # 中文数字 + 、．.
    m = re.match(r'^([一二三四五六七八九十百千]+[、．.])', text)
    if m:
        return ('chinese_dot', m.group(1))

    # 全角括号 + 中文数字
    m = re.match(r'^（([一二三四五六七八九十百千]+)）', text)
    if m:
        return ('chinese_bracket_full', f'（{m.group(1)}）')

    # 半角括号 + 中文数字
    m = re.match(r'^\(([一二三四五六七八九十百千]+)\)', text)
    if m:
        return ('chinese_bracket_half', f'({m.group(1)})')

    # 第X章/条/节/项/款
    m = re.match(r'^(第[一二三四五六七八九十百千\d]+[章条节项款])', text)
    if m:
        return ('chapter', m.group(1))

    # 阿拉伯数字 + ．.、)）
    m = re.match(r'^(\d+[．.、)）])', text)
    if m:
        return ('arabic_dot', m.group(1))

    # 全角括号 + 阿拉伯数字
    m = re.match(r'^（(\d+)）', text)
    if m:
        return ('arabic_bracket_full', f'（{m.group(1)}）')

    # 半角括号 + 阿拉伯数字
    m = re.match(r'^\((\d+)\)', text)
    if m:
        return ('arabic_bracket_half', f'({m.group(1)})')

    return (None, None)


# ═══════════════════════════════════════════════════════════════════
# 标题候选判断
# ═══════════════════════════════════════════════════════════════════

def is_title_candidate(text):
    """
    判断一段文本是否为标题候选。

    条件：
    1. 非空
    2. 不含 。！？
    3. 长度 ≤ MAX_TITLE_LENGTH
    4. 不以 ：结尾

    注：传入 text 应为段落的第一行（不含换行后的正文）。
    """
    text = text.strip()
    if not text:
        return False
    if re.search(r'[。！？]', text):
        return False
    if len(text) > MAX_TITLE_LENGTH:
        return False
    if text.endswith('：'):
        return False
    return True


# ═══════════════════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════════════════

def _is_toc_paragraph(para):
    """判断是否为 Word 自动生成的目录段落。"""
    style = para.style
    if style and style.name and style.name.startswith('TOC'):
        return True
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            val = pStyle.get(qn('w:val'))
            if val and val.startswith('TOC'):
                return True
    return False


def find_title_paragraph(doc):
    """找文档标题：第一段不含 。！？ 的文本段落。"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _is_toc_paragraph(para):
            continue
        if not re.search(r'[。！？]', text):
            return para, text
        else:
            return None, None
    return None, None


# ═══════════════════════════════════════════════════════════════════
# 标题候选收集
# ═══════════════════════════════════════════════════════════════════

class TitleCandidate:
    """标题候选，存储段落信息和分配结果。"""
    __slots__ = ('para', 'text', 'style_key', 'prefix', 'level', 'line_idx', 'run_indices')

    def __init__(self, para, text, style_key, prefix, line_idx=0, run_indices=None):
        self.para = para
        self.text = text
        self.style_key = style_key
        self.prefix = prefix
        self.level = None  # 由递归算法分配
        self.line_idx = line_idx  # 段落内行号（0=第一行）
        self.run_indices = run_indices or []  # 该行包含的 run 索引列表


def _group_runs_by_line(para):
    """
    将段落中的 runs 按行分组。

    返回 [(line_text, [run_index, ...]), ...]
    每个 tuple 对应一行：文本内容 + 该行包含的 run 索引列表。
    """
    lines = []
    current_runs = []
    current_text = []

    for i, run in enumerate(para.runs):
        has_br = run._element.find(qn('w:br')) is not None
        if has_br:
            if current_text:
                lines.append((''.join(current_text), current_runs))
            current_runs = []
            current_text = []
            continue

        text = run.text
        while '\n' in text:
            idx = text.index('\n')
            current_text.append(text[:idx])
            current_runs.append(i)
            lines.append((''.join(current_text), current_runs))
            current_runs = []
            current_text = []
            text = text[idx + 1:]

        if text:
            current_text.append(text)
            current_runs.append(i)

    if current_text:
        lines.append((''.join(current_text), current_runs))

    return lines


def _get_paragraph_numbering(para):
    """
    检测段落是否有 Word 自动编号。

    返回风格标识字符串（如 'auto_3_0' = numId=3, ilvl=0），
    无自动编号时返回 None。
    numId=0 表示"无编号"，也返回 None。
    """
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None

    numId_el = numPr.find(qn('w:numId'))
    if numId_el is None:
        return None
    num_id = numId_el.get(qn('w:val'))
    if num_id is None or num_id == '0':
        return None

    ilvl_el = numPr.find(qn('w:ilvl'))
    ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
    return f'auto_{num_id}_{ilvl}'


# 中文数字映射
CHINESE_NUMS = ['', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']


def _num_to_chinese(n):
    """将整数转换为中文数字（支持 1~99）。"""
    if 1 <= n <= 20:
        return CHINESE_NUMS[n]
    if n < 100:
        tens = n // 10
        units = n % 10
        if units == 0:
            return CHINESE_NUMS[tens] + '十'
        return CHINESE_NUMS[tens] + '十' + CHINESE_NUMS[units]
    return str(n)


ROMAN_MAP = [
    (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
    (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
    (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I'),
]


def _num_to_roman(n, upper=True):
    """将整数转换为罗马数字。"""
    result = ''
    for val, sym in ROMAN_MAP:
        while n >= val:
            result += sym
            n -= val
    return result if upper else result.lower()


def _format_number(n, fmt):
    """按格式将整数编号转为字符串。"""
    if fmt == 'decimal':
        return str(n)
    if fmt == 'chineseCounting':
        return _num_to_chinese(n)
    if fmt == 'lowerLetter':
        return chr(ord('a') + n - 1) if 1 <= n <= 26 else str(n)
    if fmt == 'upperLetter':
        return chr(ord('A') + n - 1) if 1 <= n <= 26 else str(n)
    if fmt == 'lowerRoman':
        return _num_to_roman(n, upper=False)
    if fmt == 'upperRoman':
        return _num_to_roman(n, upper=True)
    if fmt == 'ordinal':
        suffix = 'th' if 11 <= n % 100 <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')
        return str(n) + suffix
    return str(n)


def _convert_auto_numbering_to_text(doc):
    """
    预处理：将所有自动编号转为纯文本。

    解析文档的 numbering 定义，遍历所有段落：
    - 对有 w:numPr 且 numId>0 的段落，计算渲染后的编号前缀
    - 将编号前缀插入到段落第一个 run 的开头
    - 移除 w:numPr 元素（去除自动编号依赖）
    """
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return

    num_el = numbering_part._element
    if num_el is None:
        return

    # 解析 abstractNum 定义：abstractNumId → {ilvl: {start, numFmt, lvlText}}
    abstracts = {}
    for ab in num_el.findall(qn('w:abstractNum')):
        aid = ab.get(qn('w:abstractNumId'))
        if aid is None:
            continue
        levels = {}
        for lvl in ab.findall(qn('w:lvl')):
            ilvl = lvl.get(qn('w:ilvl'))
            start_el = lvl.find(qn('w:start'))
            fmt_el = lvl.find(qn('w:numFmt'))
            text_el = lvl.find(qn('w:lvlText'))
            start_val = int(start_el.get(qn('w:val'))) if start_el is not None and start_el.get(qn('w:val')) is not None else 1
            fmt_val = fmt_el.get(qn('w:val')) if fmt_el is not None else 'decimal'
            text_val = text_el.get(qn('w:val')) if text_el is not None else '%1.'
            levels[ilvl] = {
                'start': start_val,
                'numFmt': fmt_val,
                'lvlText': text_val,
            }
        abstracts[aid] = levels

    # 解析 num 定义：numId → abstractNumId
    num_map = {}
    for n in num_el.findall(qn('w:num')):
        nid = n.get(qn('w:numId'))
        if nid is None or nid == '0':
            continue
        abid_el = n.find(qn('w:abstractNumId'))
        if abid_el is not None:
            abid_val = abid_el.get(qn('w:val'))
            if abid_val is not None:
                num_map[nid] = abid_val

    # 跟踪每个 numId+ilvl 的当前编号值
    counters = {}

    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue

        numId_el = numPr.find(qn('w:numId'))
        if numId_el is None:
            continue
        num_id = numId_el.get(qn('w:val'))
        if num_id is None or num_id == '0':
            continue

        ilvl_el = numPr.find(qn('w:ilvl'))
        ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'

        # 查 numbering 定义
        abs_id = num_map.get(num_id)
        if abs_id is None:
            continue
        levels = abstracts.get(abs_id)
        if levels is None:
            continue
        lvl_def = levels.get(ilvl)
        if lvl_def is None:
            continue

        # 递增计数器
        key = (num_id, ilvl)
        if key not in counters:
            counters[key] = lvl_def['start']
        else:
            counters[key] += 1
        num_val = counters[key]

        # 格式化编号
        fmt = lvl_def['numFmt']
        formatted = _format_number(num_val, fmt)

        # 生成前缀（替换 %1、%2 等占位符）
        prefix = lvl_def['lvlText']
        prefix = prefix.replace('%1', formatted)
        prefix = re.sub(r'%\d+', '', prefix)

        # 插入到段落第一个 run 的开头
        if para.runs:
            first_run = para.runs[0]
            first_run.text = prefix + first_run.text
        else:
            # 没有 run 就创建一个
            from docx.oxml import OxmlElement
            r_elem = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.text = prefix
            t_elem.set(qn('xml:space'), 'preserve')
            r_elem.append(t_elem)
            # 在 pPr 后面插入
            after = pPr
            for child in para._element:
                if child.tag == qn('w:pPr'):
                    continue
                after = child
                break
            para._element.insert(list(para._element).index(after), r_elem)

        # 移除自动编号（w:numPr）
        pPr.remove(numPr)


def _remove_empty_paragraphs(doc):
    """移除所有空段落（正文中的空行），保留含图片的段落。"""
    image_paras = _find_image_paragraphs(doc)
    removed = 0

    for para in list(doc.paragraphs):
        if para._element in image_paras:
            continue
        if not para.text.strip():
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                # 跳过包含分节符的段落（文档结构边界）
                if para._element.find(qn('w:sectPr')) is not None:
                    continue
            para._element.getparent().remove(para._element)
            removed += 1

    if removed > 0:
        print(f"  已移除 {removed} 个空行。")


def _format_tables(doc):
    """统一表格样式：单实线边框，表头灰色背景。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsdecls
    from lxml import etree

    for table in doc.tables:
        # ── 设置表格边框：单实线 ──
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)

        # 移除旧边框定义
        old_borders = tblPr.find(qn('w:tblBorders'))
        if old_borders is not None:
            tblPr.remove(old_borders)

        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')       # 0.5pt
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            borders.append(edge_el)
        tblPr.append(borders)

        # ── 表头（第一行）灰色背景 ──
        if table.rows:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._tc.insert(0, tcPr)
                # 移除旧 shading
                old_shd = tcPr.find(qn('w:shd'))
                if old_shd is not None:
                    tcPr.remove(old_shd)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9D9D9')  # 浅灰
                tcPr.append(shd)

    print(f"  已格式化 {len(doc.tables)} 个表格。")


def _find_image_paragraphs(doc):
    """找出所有包含图片的段落，返回这些段落 element 的集合。"""
    image_paras = set()
    for para in doc.paragraphs:
        drawing = para._element.findall('.//' + qn('w:drawing'))
        pict = para._element.findall('.//' + qn('w:pict'))
        if drawing or pict:
            image_paras.add(para._element)
    return image_paras


def collect_candidates(doc, title_para):
    """
    收集文档中所有标题候选（按行检测，排除 TOC 和图片题注）。
    """
    # 预扫描：找出所有图片段落
    image_paras = _find_image_paragraphs(doc)
    all_paras = list(doc.paragraphs)

    # 找出需要排除的图片题注段落：
    # 单行 + 与图片段落相邻（不论是否有自动编号）
    caption_paras = set()
    for i, para in enumerate(all_paras):
        if para._element in image_paras:
            if i > 0:
                prev = all_paras[i - 1]
                text = prev.text.strip()
                if text:
                    lines = text.split('\n')
                    if len(lines) == 1:
                        caption_paras.add(prev._element)
            if i + 1 < len(all_paras):
                nxt = all_paras[i + 1]
                text = nxt.text.strip()
                if text:
                    lines = text.split('\n')
                    if len(lines) == 1:
                        caption_paras.add(nxt._element)

    candidates = []
    for para in doc.paragraphs:
        full_text = para.text.strip()
        if not full_text:
            continue
        if _is_toc_paragraph(para):
            continue
        if title_para is not None and para._element is title_para._element:
            continue

        # 按行分组
        line_groups = _group_runs_by_line(para)

        for line_idx, (line_text, run_idxs) in enumerate(line_groups):
            text = line_text.strip()
            if not text:
                continue

            # 跳过图片题注（单行 + 无编号 + 与图片相邻）
            if para._element in caption_paras and len(line_groups) == 1:
                continue

            if not is_title_candidate(text):
                continue

            style_key, prefix = classify_numbering(text)

            # 如果纯文本无编号，检查是否有 Word 自动编号（仅限预处理未覆盖的情况）
            if style_key is None:
                num_key = _get_paragraph_numbering(para)
                if num_key:
                    style_key = num_key
                    prefix = f'[{num_key}]'

            candidates.append(TitleCandidate(
                para, text, style_key, prefix,
                line_idx=line_idx, run_indices=run_idxs
            ))

    return candidates


# ═══════════════════════════════════════════════════════════════════
# 递归区间层级分配（v4 核心）
# ═══════════════════════════════════════════════════════════════════

def _assign_level(candidates, start, end, parent_style, level_idx):
    """
    在 [start, end) 区间内递归分配层级。

    参数：
        candidates   — 标题候选列表
        start, end   — 当前区间范围
        parent_style — 父层级的编号风格（用于排除）
        level_idx    — 当前要分配的层级索引（0=h1, 1=h2, 2=h3, 3=h4）

    逻辑：
        1. 无编号标题 → 当前层级
        2. 第一个有编号的标题 → 当前层级的风格
        3. 区间内同风格 → 同层级
        4. 在有编号的标题之间递归下一层
    """
    if start >= end:
        return

    # 超过 h4 层级 → 全部归入 h4（最深层级）
    if level_idx >= 4:
        for i in range(start, end):
            if candidates[i].level is None:
                candidates[i].level = 'h4'
        return

    # 在当前区间找第一个有编号的候选（不同父风格）
    level_style = None
    split_indices = []

    for i in range(start, end):
        c = candidates[i]
        if c.level is not None:
            continue

        if c.style_key and c.style_key != parent_style:
            level_style = c.style_key
            break

    if level_style is None:
        # 区间内没有有编号的标题 → 全部设为当前层级
        for i in range(start, end):
            if candidates[i].level is None:
                candidates[i].level = LEVEL_NAMES[level_idx]
        return

    # 遍历区间：分配当前层级
    for i in range(start, end):
        c = candidates[i]
        if c.level is not None:
            continue

        if c.style_key == level_style:
            c.level = LEVEL_NAMES[level_idx]
            split_indices.append(i)
        elif c.style_key is None:
            # 无编号标题 → 当前层级
            c.level = LEVEL_NAMES[level_idx]
        # 其他有编号风格 → 暂不处理（留给更深递归）

    # 在 split_indices 之间递归下一层
    # 同时处理第一个 split 之前的区间（可能有其他编号风格）
    prev = start
    for idx in split_indices:
        if prev < idx:
            _assign_level(candidates, prev, idx, level_style, level_idx + 1)
        prev = idx + 1
    if prev < end:
        _assign_level(candidates, prev, end, level_style, level_idx + 1)


def assign_all_levels(candidates):
    """
    对所有标题候选分配层级（h1~h4）。

    先从全局找 h1 风格，再在每个 h1 之间递归找 h2/h3/h4。
    """
    if not candidates:
        return

    # 全局找 h1：第一个有编号的候选
    h1_style = None
    h1_indices = []

    for c in candidates:
        if c.style_key:
            h1_style = c.style_key
            break

    if h1_style is None:
        # 全文档都没有有编号的标题 → 全部 h1
        for c in candidates:
            c.level = 'h1'
        return

    # 分配 h1
    for i, c in enumerate(candidates):
        if c.style_key == h1_style:
            c.level = 'h1'
            h1_indices.append(i)

    # 在每个 h1 之间递归分配 h2/h3/h4
    for j in range(len(h1_indices)):
        sub_start = h1_indices[j] + 1
        if j + 1 < len(h1_indices):
            sub_end = h1_indices[j + 1]
        else:
            sub_end = len(candidates)
        _assign_level(candidates, sub_start, sub_end, h1_style, 1)


# ═══════════════════════════════════════════════════════════════════
# 编号重写（按规范优化标题编号）
# ═══════════════════════════════════════════════════════════════════

def _replace_prefix_in_para(candidate, new_prefix):
    """替换段落中的编号前缀为新格式。"""
    para = candidate.para
    old_prefix = candidate.prefix
    run_indices = candidate.run_indices

    if not run_indices:
        return

    first_run_idx = run_indices[0]
    if first_run_idx >= len(para.runs):
        return

    run = para.runs[first_run_idx]

    if old_prefix:
        # 替换旧前缀
        if run.text.startswith(old_prefix):
            run.text = new_prefix + run.text[len(old_prefix):]
        elif old_prefix in run.text:
            run.text = run.text.replace(old_prefix, new_prefix, 1)
    else:
        # 无旧前缀，直接插入
        run.text = new_prefix + run.text


def rewrite_numbering(candidates):
    """
    按照规范重新编写标题编号（h1~h4）：

    h1 → 一、二、三、……（中文数字 + 顿号）
    h2 → （一）（二）（三）……（括号中文数字）
    h3 → 1. 2. 3. ……（阿拉伯数字 + 点）
    h4 → （1）（2）（3）……（括号阿拉伯数字）

    计数器按层级重置：h2 在每个 h1 区间内重新计数，
    h3 在每个 h2 内，h4 在每个 h3 内。
    """
    h1_counter = 0
    h2_counter = 0
    h3_counter = 0
    h4_counter = 0

    for c in candidates:
        if c.level == 'h1':
            h1_counter += 1
            h2_counter = 0
            h3_counter = 0
            h4_counter = 0
            new_prefix = _num_to_chinese(h1_counter) + '、'
        elif c.level == 'h2':
            h2_counter += 1
            h3_counter = 0
            h4_counter = 0
            new_prefix = '（' + _num_to_chinese(h2_counter) + '）'
        elif c.level == 'h3':
            h3_counter += 1
            h4_counter = 0
            new_prefix = str(h3_counter) + '.'
        elif c.level == 'h4':
            h4_counter += 1
            new_prefix = '（' + str(h4_counter) + '）'
        else:
            continue

        _replace_prefix_in_para(c, new_prefix)


# ═══════════════════════════════════════════════════════════════════
# 字体设置
# ═══════════════════════════════════════════════════════════════════

def _set_run_font(run, font_name, font_size, bold=False):
    """设置单个 run 的字体属性。"""
    run.font.size = font_size
    run.font.bold = bold

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

    # 显式移除加粗（确保覆盖原文样式）
    b_elem = rPr.find(qn('w:b'))
    if not bold:
        if b_elem is not None:
            rPr.remove(b_elem)
        # 同时移除 w:bCs（复杂脚本加粗）
        bCs = rPr.find(qn('w:bCs'))
        if bCs is not None:
            rPr.remove(bCs)
    else:
        if b_elem is None:
            b_elem = OxmlElement('w:b')
            rPr.append(b_elem)


def _set_paragraph_format(paragraph, alignment, line_spacing=None, first_line_indent=None):
    """设置段落格式。"""
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format

    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = line_spacing

    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def _apply_run_level(paragraph, level, run_indices):
    """仅应用 run 级格式（字体、字号、粗体），不修改段落级格式。"""
    if level not in FONT_NAMES:
        return
    font_name = FONT_NAMES[level]
    font_size = FONT_SIZES[level]
    bold = FONT_BOLD[level]
    for idx in run_indices:
        if idx < len(paragraph.runs):
            _set_run_font(paragraph.runs[idx], font_name, font_size, bold)


def apply_style(paragraph, level, run_indices=None):
    """对段落（或指定 run）应用指定层级的完整格式。"""
    if level not in FONT_NAMES:
        return

    font_name = FONT_NAMES[level]
    font_size = FONT_SIZES[level]
    bold = FONT_BOLD[level]
    alignment = ALIGNMENT[level]

    if level == 'body':
        _set_paragraph_format(paragraph, alignment, LINE_SPACING_FIXED, FIRST_LINE_INDENT)
    else:
        _set_paragraph_format(paragraph, alignment)

    if run_indices is not None:
        # 只格式化指定 run
        for idx in run_indices:
            if idx < len(paragraph.runs):
                _set_run_font(paragraph.runs[idx], font_name, font_size, bold)
    else:
        # 格式化所有 run
        for run in paragraph.runs:
            _set_run_font(run, font_name, font_size, bold)


# ═══════════════════════════════════════════════════════════════════
# 主流程（v4 重写）
# ═══════════════════════════════════════════════════════════════════

def format_document(doc):
    """
    三遍扫描格式化文档。

    第一遍：找文档标题。
    第二遍：收集标题候选，递归分配层级。
    第三遍：逐段应用格式。
    """
    # ── 第零遍：预处理——自动编号转纯文本 ──
    _convert_auto_numbering_to_text(doc)
    print("  自动编号已转为纯文本。")

    # ── 第一遍：找标题 ──
    title_para, title_text = find_title_paragraph(doc)
    if title_para is not None:
        print(f"  识别标题：「{title_text[:40]}{'…' if len(title_text) > 40 else ''}」")

    # ── 第二遍：收集标题候选 + 递归分配层级 + 编号重写 ──
    candidates = collect_candidates(doc, title_para)
    assign_all_levels(candidates)
    rewrite_numbering(candidates)

    # 打印分配结果
    if candidates:
        level_counts = {}
        for c in candidates:
            level_counts[c.level] = level_counts.get(c.level, 0) + 1
        # 过滤掉 None（未分配），防止 sorted 报错
        sorted_items = sorted((lvl, cnt) for lvl, cnt in level_counts.items() if lvl is not None)
        parts = [f"    {lvl}×{cnt}" for lvl, cnt in sorted_items]
        if parts:
            print(f"  标题分配：{' '.join(parts)}")
        unassigned = [c for c in candidates if c.level is None]
        if unassigned:
            print(f"  ⚠️ 有 {len(unassigned)} 个候选未分配层级：")
            for c in unassigned:
                print(f"    「{c.text[:40]}」")
        for c in candidates:
            prefix_info = f"[{c.prefix}]" if c.prefix else "[无编号]"
            level_str = c.level if c.level else 'None'
            print(f"    {level_str:4s} {prefix_info:12s} {c.text[:40]}")

    # 构建段落→层级查找表
    # 对于多行段落，存储 {(para._element, line_idx): level}
    para_level = {}  # (element, line_idx) → level
    para_lines = {}  # element → [(line_text, run_indices, level_or_None), ...]
    for c in candidates:
        para_level[(c.para._element, c.line_idx)] = c.level

    # ── 第三遍：逐段应用格式 ──
    title_applied = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if _is_toc_paragraph(para):
            continue

        # 判断是否有多行标题候选
        line_groups = _group_runs_by_line(para)
        has_any_heading = any((para._element, i) in para_level for i in range(len(line_groups)))

        if has_any_heading:
            # 找到第一个标题行的层级，用于段落级格式
            first_heading_level = None
            for i in range(len(line_groups)):
                lvl = para_level.get((para._element, i))
                if lvl:
                    first_heading_level = lvl
                    break

            # 段落级格式：只用第一个标题行的层级（一次设置，后续不覆盖）
            if first_heading_level:
                apply_style(para, first_heading_level, run_indices=[])

            # 逐行应用 run 级格式
            # 第一个标题行已在 apply_style(ri=[]) 中设置过段落级格式，
            # 后续行（含标题行和正文行）只改 run 级格式，不覆盖段落级
            first_heading_applied = False
            for line_idx, (line_text, run_idxs) in enumerate(line_groups):
                level = para_level.get((para._element, line_idx))
                if level:
                    if not first_heading_applied:
                        apply_style(para, level, run_indices=run_idxs)
                        first_heading_applied = True
                    else:
                        _apply_run_level(para, level, run_idxs)
                else:
                    _apply_run_level(para, 'body', run_idxs)
        elif not title_applied and title_para is not None and para._element is title_para._element:
            level = 'title'
            title_applied = True
            apply_style(para, level)
        else:
            apply_style(para, 'body')

    if not title_applied:
        print("  未找到明确的标题段落。")
    else:
        print("  标题格式已应用。")

    # ── 第四遍：收尾——移除空行 + 表格统一样式 ──
    _remove_empty_paragraphs(doc)
    _format_tables(doc)


def main():
    if len(sys.argv) < 2:
        print("用法：format_word.exe <输入文件.docx> [输出文件.docx]")
        print("示例：format_word.exe 报告.docx")
        print("       format_word.exe 报告.docx 已格式化_报告.docx")
        input("\n按 Enter 键退出……")
        sys.exit(1)

    input_path = sys.argv[1]

    if not os.path.isfile(input_path):
        print(f"错误：文件不存在 — {input_path}")
        input("\n按 Enter 键退出……")
        sys.exit(1)

    if not input_path.lower().endswith('.docx'):
        print(f"错误：仅支持 .docx 格式 — {input_path}")
        input("\n按 Enter 键退出……")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        base, ext = os.path.splitext(input_path)
        output_path = f"{base}_格式化{ext}"

    print(f"📖 读取：{input_path}")
    doc = Document(input_path)

    print("🔧 格式化中……")
    format_document(doc)

    print(f"💾 保存：{output_path}")
    doc.save(output_path)

    print("✅ 完成！")
    try:
        input("\n按 Enter 键退出……")
    except EOFError:
        pass


if __name__ == '__main__':
    main()